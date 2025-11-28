import streamlit as st
import tabula
from docx import Document
import tempfile
import os

st.set_page_config(page_title="PDF → Word (Giữ bảng)", layout="wide")
st.title("📄 Chuyển PDF → Word (Giữ nguyên dữ liệu bảng)")

st.write("Ứng dụng chuyển PDF sang Word và giữ nguyên dữ liệu bảng (table).")


# =====================================================================
# FUNCTION: PDF → LIST OF DATAFRAMES
# =====================================================================
def extract_tables(pdf_file):
    dfs = tabula.read_pdf(
        pdf_file,
        pages="all",
        multiple_tables=True,
        stream=True  # đọc dạng dòng, tránh gãy bảng
    )
    return dfs


# =====================================================================
# FUNCTION: WRITE TABLES TO WORD
# =====================================================================
def create_word_from_tables(dataframes):
    doc = Document()

    for idx, df in enumerate(dataframes):
        doc.add_heading(f"Bảng {idx + 1}", level=2)

        # tạo bảng Word với số cột tương ứng
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells

        # header
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        # data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        doc.add_paragraph("")  # thêm khoảng trắng giữa các bảng

    return doc


# =====================================================================
# STREAMLIT UI
# =====================================================================

uploaded_file = st.file_uploader("📤 Chọn file PDF", type="pdf")

if uploaded_file:
    st.success("PDF đã tải lên!")

    # lưu file tạm
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf.write(uploaded_file.read())
    temp_pdf.close()

    if st.button("🔍 Trích bảng từ PDF"):
        with st.spinner("Đang phân tích và trích bảng..."):
            tables = extract_tables(temp_pdf.name)

        if not tables:
            st.error("❌ Không tìm thấy bảng nào trong PDF!")
        else:
            st.success(f"✔ Tìm thấy {len(tables)} bảng trong PDF!")
            
            # hiển thị preview
            for i, df in enumerate(tables):
                st.subheader(f"Bảng {i+1}")
                st.dataframe(df)

            # tạo Word
            word_doc = create_word_from_tables(tables)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            word_doc.save(output_path)

            # download button
            with open(output_path, "rb") as f:
                st.download_button(
                    label="📥 Tải file Word",
                    data=f,
                    file_name="output_tables.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # cleanup
    os.unlink(temp_pdf.name)
